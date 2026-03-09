"""Generate synthetic test documents with known DREF field values for e2e testing."""
import base64
import io

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


# Ground truth: the field values we embed in test documents
FLOOD_REPORT_FIELDS = {
    "operation_overview.country": "Sudan",
    "operation_overview.disaster_type": "Flood",
    "operation_overview.national_society": "Sudan Red Crescent Society",
    "event_detail.total_affected_population": "12500",
    "event_detail.what_happened": "Severe flooding in Blue Nile State displaced thousands of families",
    "event_detail.people_in_need": "8000",
}

# A second document with SOME overlapping fields but DIFFERENT values (for conflict testing)
CONFLICTING_REPORT_FIELDS = {
    "operation_overview.country": "Sudan",  # same
    "operation_overview.disaster_type": "Flood",  # same
    "event_detail.total_affected_population": "15000",  # DIFFERENT
    "event_detail.people_in_need": "10000",  # DIFFERENT
    "event_detail.what_happened": "Flash flooding across Blue Nile and Sennar states affected rural communities",
}


def generate_flood_report_pdf() -> tuple[str, str]:
    """Generate a PDF with known DREF flood report data.

    Returns:
        Tuple of (base64_data, filename)
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Sudan Flood Situation Report")

    y -= 40
    c.setFont("Helvetica", 11)
    lines = [
        "Country: Sudan",
        "Disaster Type: Flood",
        "National Society: Sudan Red Crescent Society",
        "",
        "Event Description:",
        "Severe flooding in Blue Nile State displaced thousands of families.",
        "The floods began on March 1, 2026 following heavy seasonal rains.",
        "",
        "Affected Population: 12,500 people",
        "People in Need: 8,000 people",
        "",
        "The Sudan Red Crescent Society has deployed emergency response teams",
        "to the affected areas and is coordinating with local authorities.",
    ]

    for line in lines:
        c.drawString(50, y, line)
        y -= 18

    c.save()
    b64 = base64.b64encode(buf.getvalue()).decode()
    return b64, "flood_situation_report.pdf"


def generate_flood_report_docx() -> tuple[str, str]:
    """Generate a DOCX with known DREF flood report data.

    Returns:
        Tuple of (base64_data, filename)
    """
    doc = DocxDocument()
    doc.add_heading("Sudan Flood Situation Report", level=1)
    doc.add_paragraph(
        "Severe flooding in Blue Nile State displaced thousands of families. "
        "The floods began on March 1, 2026 following heavy seasonal rains."
    )
    doc.add_paragraph("The Sudan Red Crescent Society has been responding since day one.")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    data = [
        ("Country", "Sudan"),
        ("Disaster Type", "Flood"),
        ("National Society", "Sudan Red Crescent Society"),
        ("Total Affected Population", "12,500"),
        ("People in Need", "8,000"),
    ]
    for i, (label, value) in enumerate(data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()
    return b64, "flood_situation_report.docx"


def generate_conflicting_assessment_docx() -> tuple[str, str]:
    """Generate a DOCX with DIFFERENT numbers for the same fields (for conflict testing).

    Returns:
        Tuple of (base64_data, filename)
    """
    doc = DocxDocument()
    doc.add_heading("Needs Assessment Report - Blue Nile & Sennar", level=1)
    doc.add_paragraph(
        "Flash flooding across Blue Nile and Sennar states affected rural communities. "
        "An estimated 15,000 people have been affected and 10,000 are in immediate need."
    )

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    data = [
        ("Country", "Sudan"),
        ("Total Affected", "15,000"),
        ("People in Need", "10,000"),
    ]
    for i, (label, value) in enumerate(data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()
    return b64, "needs_assessment.docx"


def generate_test_image_png() -> tuple[str, str]:
    """Generate a minimal PNG test image (1x1 red pixel).

    Returns:
        Tuple of (base64_data, filename)
    """
    import struct
    import zlib

    def create_minimal_png():
        signature = b'\x89PNG\r\n\x1a\n'
        ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xffffffff
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
        raw = zlib.compress(b'\x00\xff\x00\x00')
        idat_crc = zlib.crc32(b'IDAT' + raw) & 0xffffffff
        idat = struct.pack('>I', len(raw)) + b'IDAT' + raw + struct.pack('>I', idat_crc)
        iend_crc = zlib.crc32(b'IEND') & 0xffffffff
        iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
        return signature + ihdr + idat + iend

    png_bytes = create_minimal_png()
    b64 = base64.b64encode(png_bytes).decode()
    return b64, "damage_photo.png"
