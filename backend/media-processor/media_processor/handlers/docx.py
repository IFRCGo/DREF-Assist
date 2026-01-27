"""DOCX file handler - text and image extraction."""
import base64
import io

from docx import Document
from docx.oxml.ns import qn

from media_processor.models import FileInput, SourceResult


class DOCXHandler:
    """Handler for DOCX files."""

    MAX_SIZE_BYTES = 25 * 1024 * 1024  # 25MB
    DRAWING_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    BLIP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    EMBED_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

    def process(self, file_input: FileInput, source_index: int) -> SourceResult:
        """Process a DOCX file.

        Pipeline:
        1. Iterate paragraphs and tables in document order
        2. Insert [IMAGE_N] placeholder when InlineShape encountered
        3. Extract embedded images

        Args:
            file_input: The file input containing base64 data
            source_index: Index for namespacing (DOCX_N_IMAGE_M)

        Returns:
            SourceResult with text and images
        """
        try:
            # Decode base64
            docx_bytes = base64.b64decode(file_input.data)

            # Check size limit
            if len(docx_bytes) > self.MAX_SIZE_BYTES:
                return SourceResult(
                    filename=file_input.filename,
                    source_type="docx",
                    error=f"File exceeds size limit of 25MB",
                    images={},
                )

            # Open document
            try:
                doc = Document(io.BytesIO(docx_bytes))
            except Exception as e:
                return SourceResult(
                    filename=file_input.filename,
                    source_type="docx",
                    error=f"Failed to parse DOCX: file corrupted or invalid",
                    images={},
                )

            text_parts = []
            images: dict[str, str] = {}
            image_counter = 0
            table_counter = 0

            # Process paragraphs
            for para in doc.paragraphs:
                para_text = para.text.strip()

                # Check for inline images in this paragraph
                inline_shapes = para._element.xpath('.//a:blip', namespaces={
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                })

                for blip in inline_shapes:
                    embed_attr = blip.get(qn('r:embed'))
                    if embed_attr and embed_attr in doc.part.related_parts:
                        image_counter += 1
                        image_key = f"DOCX_{source_index}_IMAGE_{image_counter}"

                        try:
                            image_part = doc.part.related_parts[embed_attr]
                            images[image_key] = base64.b64encode(image_part.blob).decode()

                            # Insert placeholder
                            if para_text:
                                para_text += f" [{image_key}]"
                            else:
                                para_text = f"[{image_key}]"
                        except Exception:
                            continue

                if para_text:
                    text_parts.append(para_text)

            # Process tables
            for table in doc.tables:
                table_counter += 1
                table_text_parts = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):
                        table_text_parts.append(" | ".join(row_cells))

                if table_text_parts:
                    text_parts.append("\n".join(table_text_parts))

            text_content = "\n\n".join(text_parts)

            notes_parts = []
            if image_counter:
                notes_parts.append(f"Extracted {image_counter} embedded images")
            if table_counter:
                notes_parts.append(f"Processed {table_counter} tables")

            return SourceResult(
                filename=file_input.filename,
                source_type="docx",
                text_content=text_content,
                images=images,
                processing_notes=". ".join(notes_parts) if notes_parts else None,
            )

        except Exception as e:
            return SourceResult(
                filename=file_input.filename,
                source_type="docx",
                error=f"Failed to process DOCX: {str(e)}",
                images={},
            )
