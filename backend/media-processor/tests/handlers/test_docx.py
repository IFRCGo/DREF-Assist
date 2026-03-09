import base64
import io
import pytest
from unittest.mock import patch, Mock, MagicMock, PropertyMock

from media_processor.handlers.docx import DOCXHandler
from media_processor.models import FileInput, FileType


class TestDOCXHandler:
    @patch("media_processor.handlers.docx.Document")
    def test_process_text_only_docx(self, mock_document_class):
        mock_doc = MagicMock()
        mock_para1 = MagicMock()
        mock_para1.text = "First paragraph."
        mock_para1._element.xpath.return_value = []  # No inline shapes

        mock_para2 = MagicMock()
        mock_para2.text = "Second paragraph."
        mock_para2._element.xpath.return_value = []

        mock_doc.paragraphs = [mock_para1, mock_para2]
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        handler = DOCXHandler()
        data = base64.b64encode(b"fake docx").decode()
        file_input = FileInput(data=data, type=FileType.DOCX, filename="test.docx")

        result = handler.process(file_input, source_index=1)

        assert result.filename == "test.docx"
        assert result.source_type == "docx"
        assert "First paragraph" in result.text_content
        assert "Second paragraph" in result.text_content
        assert result.error is None

    @patch("media_processor.handlers.docx.Document")
    @patch("media_processor.handlers.docx.qn")
    def test_process_docx_with_images(self, mock_qn, mock_document_class):
        mock_doc = MagicMock()

        # Paragraph with inline image
        mock_para = MagicMock()
        mock_para.text = "Text with image."

        # Mock blip element (returned directly by xpath './/a:blip')
        mock_blip = MagicMock()
        # Mock qn to return the namespace-qualified attribute name
        mock_qn.return_value = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        mock_blip.get.return_value = "rId1"

        # The findall returns blip elements directly
        mock_para._element.findall.return_value = [mock_blip]

        # Mock image part
        mock_image_part = MagicMock()
        mock_image_part.blob = b"image bytes"
        mock_image_part.content_type = "image/png"
        mock_doc.part.related_parts = {"rId1": mock_image_part}

        mock_doc.paragraphs = [mock_para]
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        handler = DOCXHandler()
        data = base64.b64encode(b"docx with image").decode()
        file_input = FileInput(data=data, type=FileType.DOCX, filename="report.docx")

        result = handler.process(file_input, source_index=1)

        assert result.error is None
        assert len(result.images) > 0
        assert "[DOCX_1_IMAGE_" in result.text_content

    @patch("media_processor.handlers.docx.Document")
    def test_process_docx_with_tables(self, mock_document_class):
        mock_doc = MagicMock()

        mock_para = MagicMock()
        mock_para.text = "Before table."
        mock_para._element.xpath.return_value = []

        # Mock table
        mock_cell1 = MagicMock()
        mock_cell1.text = "Cell 1"
        mock_cell2 = MagicMock()
        mock_cell2.text = "Cell 2"
        mock_row = MagicMock()
        mock_row.cells = [mock_cell1, mock_cell2]
        mock_table = MagicMock()
        mock_table.rows = [mock_row]

        mock_doc.paragraphs = [mock_para]
        mock_doc.tables = [mock_table]

        # Mock document body order
        mock_doc.element.body = MagicMock()
        mock_doc.element.body.__iter__ = Mock(return_value=iter([
            mock_para._element,
            mock_table._tbl,
        ]))

        mock_document_class.return_value = mock_doc

        handler = DOCXHandler()
        data = base64.b64encode(b"docx with table").decode()
        file_input = FileInput(data=data, type=FileType.DOCX, filename="data.docx")

        result = handler.process(file_input, source_index=1)

        assert result.error is None
        # Should include table content
        assert "Cell 1" in result.text_content or "table" in result.processing_notes.lower()

    def test_reject_oversized_docx(self):
        handler = DOCXHandler()
        large_data = base64.b64encode(b"x" * (26 * 1024 * 1024)).decode()
        file_input = FileInput(data=large_data, type=FileType.DOCX, filename="huge.docx")

        result = handler.process(file_input, source_index=1)

        assert result.error is not None
        assert "25MB" in result.error

    @patch("media_processor.handlers.docx.Document")
    def test_corrupted_docx_returns_error(self, mock_document_class):
        mock_document_class.side_effect = Exception("Cannot open DOCX")

        handler = DOCXHandler()
        data = base64.b64encode(b"corrupted").decode()
        file_input = FileInput(data=data, type=FileType.DOCX, filename="bad.docx")

        result = handler.process(file_input, source_index=1)

        assert result.error is not None


def test_end_to_end_docx_with_known_content():
    """Generate a real DOCX with known disaster data and verify extraction."""
    from docx import Document
    import base64
    import io
    from media_processor.models import FileInput
    from media_processor.handlers.docx import DOCXHandler

    doc = Document()
    doc.add_heading("Sudan Flood Situation Report", level=1)
    doc.add_paragraph("The flooding in Blue Nile State affected 12,500 people. 45 people were reported deceased.")
    doc.add_paragraph("The Sudan Red Crescent Society deployed 30 volunteers.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Total Affected"
    table.cell(0, 1).text = "12,500"
    table.cell(1, 0).text = "Deceased"
    table.cell(1, 1).text = "45"
    table.cell(2, 0).text = "Disaster Type"
    table.cell(2, 1).text = "Flood"

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    handler = DOCXHandler()
    result = handler.process(
        FileInput(data=b64, type="docx", filename="sudan_sitrep.docx"),
        source_index=1,
    )

    assert result.error is None
    assert "12,500" in result.text_content
    assert "45" in result.text_content
    assert "Sudan Red Crescent" in result.text_content
    assert "Flood" in result.text_content
